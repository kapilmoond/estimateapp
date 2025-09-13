"""
Document processing module for various file types
Handles PDF, DOCX, TXT, XLSX, XLS files
"""

import io
import re
import logging
from typing import Dict, List, Any
import asyncio

import PyPDF2
from docx import Document
import pandas as pd
import openpyxl
import tiktoken

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    async def process_file(self, file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """Process a file and return extracted content and chunks"""
        try:
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text = await self.extract_pdf_text(file_content)
            elif file_type.lower() == 'docx':
                text = await self.extract_docx_text(file_content)
            elif file_type.lower() in ['txt']:
                text = await self.extract_txt_text(file_content)
            elif file_type.lower() in ['xlsx', 'xls']:
                text = await self.extract_excel_text(file_content, file_type)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Clean and normalize text
            text = self.clean_text(text)
            
            # Create chunks
            chunks = self.create_chunks(text)
            
            # Calculate word count
            word_count = len(text.split())
            
            return {
                "content": text,
                "chunks": chunks,
                "word_count": word_count,
                "chunk_count": len(chunks)
            }
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise
    
    async def extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    async def extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    async def extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            return text
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise
    
    async def extract_excel_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text from Excel file"""
        try:
            excel_file = io.BytesIO(file_content)
            
            if file_type.lower() == 'xlsx':
                # Use openpyxl for .xlsx files
                workbook = openpyxl.load_workbook(excel_file, data_only=True)
                text = ""
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"\n--- Sheet: {sheet_name} ---\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell in row:
                            if cell is not None and str(cell).strip():
                                row_text.append(str(cell).strip())
                        if row_text:
                            text += " | ".join(row_text) + "\n"
            else:
                # Use pandas for .xls files
                excel_data = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
                text = ""
                
                for sheet_name, df in excel_data.items():
                    text += f"\n--- Sheet: {sheet_name} ---\n"
                    
                    # Convert DataFrame to text
                    for _, row in df.iterrows():
                        row_text = []
                        for value in row:
                            if pd.notna(value) and str(value).strip():
                                row_text.append(str(value).strip())
                        if row_text:
                            text += " | ".join(row_text) + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting Excel text: {e}")
            raise
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create chunks from text with overlap"""
        chunks = []
        
        # Split text into sentences for better chunk boundaries
        sentences = self.split_into_sentences(text)
        
        current_chunk = ""
        current_tokens = 0
        chunk_index = 0
        start_char = 0
        
        for sentence in sentences:
            sentence_tokens = len(self.tokenizer.encode(sentence))
            
            # If adding this sentence would exceed chunk size, finalize current chunk
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                # Create chunk
                chunk_content = current_chunk.strip()
                if chunk_content:
                    chunks.append({
                        "id": f"chunk_{chunk_index}",
                        "content": chunk_content,
                        "chunkIndex": chunk_index,
                        "metadata": {
                            "startChar": start_char,
                            "endChar": start_char + len(chunk_content),
                            "tokenCount": current_tokens
                        }
                    })
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_text = self.get_overlap_text(current_chunk, self.chunk_overlap)
                current_chunk = overlap_text + " " + sentence if overlap_text else sentence
                current_tokens = len(self.tokenizer.encode(current_chunk))
                start_char += len(chunk_content) - len(overlap_text)
            else:
                # Add sentence to current chunk
                current_chunk += " " + sentence if current_chunk else sentence
                current_tokens += sentence_tokens
        
        # Add final chunk if it has content
        if current_chunk.strip():
            chunk_content = current_chunk.strip()
            chunks.append({
                "id": f"chunk_{chunk_index}",
                "content": chunk_content,
                "chunkIndex": chunk_index,
                "metadata": {
                    "startChar": start_char,
                    "endChar": start_char + len(chunk_content),
                    "tokenCount": current_tokens
                }
            })
        
        return chunks
    
    def split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting - can be improved with more sophisticated methods
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def get_overlap_text(self, text: str, max_tokens: int) -> str:
        """Get overlap text from the end of current chunk"""
        if not text or max_tokens <= 0:
            return ""
        
        words = text.split()
        overlap_words = []
        token_count = 0
        
        # Take words from the end until we reach max_tokens
        for word in reversed(words):
            word_tokens = len(self.tokenizer.encode(word))
            if token_count + word_tokens > max_tokens:
                break
            overlap_words.insert(0, word)
            token_count += word_tokens
        
        return " ".join(overlap_words)
